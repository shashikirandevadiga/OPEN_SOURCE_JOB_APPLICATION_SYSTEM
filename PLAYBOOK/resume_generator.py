#!/usr/bin/env python3
"""
Resume Generator V3 - Pixel-Perfect Template Match

CRITICAL FIXES:
- Summary: 3 lines (natural wrapping)
- Skills: 1 line per category (NO wrapping)
- Education: Year right-aligned with tab stops
- Page width: Exact margins matching template
- Spacing: Ultra-tight to fit on 1 page

Usage:
    python3 generate_resume_v3.py -i Template.md -o Template_REGENERATED.docx
"""

import argparse
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Company to location mapping
# Users should update this dictionary with their own company → location mappings
# Example format: 'Company Name': 'City, Country'
COMPANY_LOCATIONS = {
    # Add your companies and locations here
    # 'Company 1': 'City, Country',
    # 'Company 2': 'City, Country',
}


def add_underline_to_paragraph(paragraph):
    """Add bottom border (underline) to paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper over a <w:r> element)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    # Set size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')  # 10pt = 20 half-points
    rPr.append(sz)

    # Set color (blue for hyperlinks)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink


def set_calibri_font(run, size_pt, bold=False, italic=False):
    """Set font to Calibri with specific size"""
    run.font.name = 'Calibri'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Calibri')


def parse_bold_text(text):
    """Parse markdown bold (**text**) and return list of (text, is_bold) tuples"""
    parts = []
    pattern = r'\*\*(.*?)\*\*'
    last_end = 0

    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        parts.append((match.group(1), True))
        last_end = match.end()

    if last_end < len(text):
        parts.append((text[last_end:], False))

    return parts if parts else [(text, False)]


def add_mixed_format_text(paragraph, text, font_size=10):
    """Add text with mixed bold/normal formatting"""
    parts = parse_bold_text(text)
    for part_text, is_bold in parts:
        run = paragraph.add_run(part_text)
        set_calibri_font(run, font_size, bold=is_bold)


def clean_markdown(text):
    """Remove markdown formatting"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text.strip()


def parse_markdown_link(text):
    """Parse markdown [display_text](url) and return (display_text, url) or None

    Example:
        "[LinkedIn](https://linkedin.com/in/user)" -> ("LinkedIn", "https://linkedin.com/in/user")
        "plain text" -> None
    """
    match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', text.strip())
    if match:
        return (match.group(1), match.group(2))
    return None


def create_resume_from_markdown(md_file_path, output_path):
    """Generate PERFECT Resume.docx from RESUME.md matching template exactly"""

    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # CRITICAL: Set exact margins matching template
    for section in doc.sections:
        section.top_margin = Inches(0.46)     # Template: 0.46"
        section.bottom_margin = Inches(0.19)  # Template: 0.19"
        section.left_margin = Inches(0.50)    # Template: 0.50"
        section.right_margin = Inches(0.25)   # Template: 0.25"

    lines = content.split('\n')
    i = 0
    summary_lines = []

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines and horizontal rules
        if not line or line.startswith('---'):
            i += 1
            continue

        # NAME (# [YOUR_NAME])
        if line.startswith('# '):
            name = clean_markdown(line[2:])
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(0)  # Ultra-tight
            run = para.add_run(name.upper())
            set_calibri_font(run, 14, bold=True)
            i += 1
            continue

        # CONTACT LINE
        if '@' in line and '|' in line and i < 5:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(6)  # Tight spacing
            contact_parts = [part.strip() for part in line.split('|')]

            for idx, part in enumerate(contact_parts):
                # Add bullet separator (except before first item)
                if idx > 0:
                    run = para.add_run(' • ')
                    set_calibri_font(run, 10)

                # Check if this part is a markdown hyperlink [text](url)
                link_data = parse_markdown_link(part)
                if link_data:
                    # It's a markdown link - extract display text and URL
                    display_text, url = link_data
                    add_hyperlink(para, display_text, url)
                else:
                    # Add as normal text (plain text parts like location, phone, email)
                    run = para.add_run(part)
                    set_calibri_font(run, 10)

            i += 1
            continue

        # SUMMARY (multi-line paragraph - BOLD, let it wrap naturally)
        if i < 15 and not line.startswith('#') and not line.startswith('•') and len(line) > 50:
            summary_lines.append(line)
            i += 1
            if i < len(lines) and len(lines[i].strip()) > 50 and not lines[i].startswith('#'):
                continue
            if summary_lines:
                summary_text = ' '.join(summary_lines)
                # CRITICAL: Strip ** markdown markers before adding to document
                summary_text = clean_markdown(summary_text)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(6)  # Tight
                para.paragraph_format.line_spacing = 1.0  # Single spacing
                run = para.add_run(summary_text)
                set_calibri_font(run, 10, bold=True)  # BOLD summary
                summary_lines = []
            continue

        # SECTION HEADERS (## PROFESSIONAL EXPERIENCE, ## SKILLS, ## EDUCATION)
        if line.startswith('## '):
            header_text = clean_markdown(line[3:]).upper()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)  # Add gap before section
            para.paragraph_format.space_after = Pt(2)   # Small gap after underline
            para.paragraph_format.line_spacing = 1.0   # CRITICAL: 1.0 line spacing (not 1.5)
            run = para.add_run(header_text)
            set_calibri_font(run, 10, bold=True)
            add_underline_to_paragraph(para)
            i += 1
            continue

        # ROLE HEADERS (### **Product Manager** | **Company** | **Year**)
        if line.startswith('### '):
            role_line = clean_markdown(line[4:])
            parts = [p.strip() for p in role_line.split('|')]

            if len(parts) >= 3:
                role = parts[0]
                company = parts[1]
                year = parts[2]

                # Create role/company/year line
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(6)   # Tight
                para.paragraph_format.space_after = Pt(0)    # Ultra-tight
                para.paragraph_format.line_spacing = 1.0     # CRITICAL: 1.0 line spacing (not 1.5)

                # Add tab stop at right edge (7.75") for right-aligned year
                tab_stops = para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.75), WD_TAB_ALIGNMENT.RIGHT)

                # Company | Role (bold)
                run1 = para.add_run(f"{company} | {role}")
                set_calibri_font(run1, 10, bold=True)

                # Tab + Year
                para.add_run('\t')
                run2 = para.add_run(year)
                set_calibri_font(run2, 10)

                i += 1

                # Check for location - skip blank lines first
                location_found = False
                temp_i = i
                while temp_i < len(lines) and not lines[temp_i].strip():
                    temp_i += 1  # Skip blank lines

                if temp_i < len(lines):
                    next_line = lines[temp_i].strip()
                    # Check if it's a location (short line with comma, not a bullet)
                    if len(next_line) < 30 and ',' in next_line and not next_line.startswith('•') and not next_line.startswith('**'):
                        loc_para = doc.add_paragraph()
                        loc_para.paragraph_format.space_after = Pt(2)
                        loc_para.paragraph_format.line_spacing = 1.0  # CRITICAL: 1.0 line spacing
                        run = loc_para.add_run(next_line)
                        set_calibri_font(run, 9, italic=True)
                        i = temp_i + 1  # Skip past location line
                        location_found = True

                if not location_found:
                    # Try to get from mapping as fallback
                    location = COMPANY_LOCATIONS.get(company, None)
                    if location:
                        loc_para = doc.add_paragraph()
                        loc_para.paragraph_format.space_after = Pt(2)
                        loc_para.paragraph_format.line_spacing = 1.0  # CRITICAL: 1.0 line spacing
                        run = loc_para.add_run(location)
                        set_calibri_font(run, 9, italic=True)

                continue

        # BULLETS (• Text or - Text)
        if line.startswith('• ') or line.startswith('- '):
            bullet_text = line[2:].strip()

            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)  # RESTORE indent (template has this!)
            para.paragraph_format.space_after = Pt(0)  # NO spacing between bullets
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            add_mixed_format_text(para, bullet_text, 10)

            i += 1
            continue

        # SKILLS SECTION - CRITICAL: Create ONE paragraph with line breaks
        if ':' in line and not line.startswith('#') and len(line.split(':')) == 2:
            # Check if this is a skills line
            parts = line.split(':', 1)
            category = clean_markdown(parts[0].strip())
            skills = clean_markdown(parts[1].strip())

            # Only process if it looks like skills (short category, longer skills list)
            if len(category) < 50 and len(skills) > 20:
                # Check if this is the FIRST skills line - create paragraph
                # Look ahead to collect ALL skills lines
                all_skills_lines = []
                j = i
                while j < len(lines):
                    current_line = lines[j].strip()
                    if ':' in current_line and not current_line.startswith('#'):
                        curr_parts = current_line.split(':', 1)
                        curr_cat = clean_markdown(curr_parts[0].strip())
                        curr_skills = clean_markdown(curr_parts[1].strip())
                        if len(curr_cat) < 50 and len(curr_skills) > 20:
                            all_skills_lines.append((curr_cat, curr_skills))
                            j += 1
                        else:
                            break
                    elif current_line.startswith('---') or current_line.startswith('##'):
                        break
                    else:
                        j += 1

                # Create ONE paragraph with all skills (line breaks between)
                if all_skills_lines:
                    para = doc.add_paragraph()
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.0  # CRITICAL: 1.0 line spacing (not 1.5)

                    for idx, (cat, skl) in enumerate(all_skills_lines):
                        # Category (bold)
                        run_cat = para.add_run(cat + ':')
                        set_calibri_font(run_cat, 10, bold=True)

                        # Skills (normal)
                        run_skills = para.add_run(' ' + skl)
                        set_calibri_font(run_skills, 10)

                        # Add line break EXCEPT after last skill
                        if idx < len(all_skills_lines) - 1:
                            para.add_run('\n')

                    # Skip ahead past all skills lines
                    i = j
                    continue

                i += 1
                continue

        # EDUCATION (University | Degree | Year)
        if '|' in line and not line.startswith('#') and ('University' in line or 'GPA' in line):
            parts = [p.strip() for p in line.split('|')]

            if len(parts) >= 3:
                # Extract year from last part (might have ** markers)
                year_part = clean_markdown(parts[-1])
                left_parts = parts[:-1]
                left_text = ' | '.join(left_parts)

                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(0)  # Tight
                para.paragraph_format.line_spacing = 1.0   # CRITICAL: 1.0 line spacing (not 1.5)

                # Add tab stop at right edge (7.75") for right-aligned year
                tab_stops = para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.75), WD_TAB_ALIGNMENT.RIGHT)

                # University | Degree | GPA (mixed bold)
                add_mixed_format_text(para, left_text, 10)

                # Tab + Year (right-aligned)
                para.add_run('\t')
                run_year = para.add_run(year_part)
                set_calibri_font(run_year, 10)

                i += 1
                continue

        # DEFAULT: Skip unrecognized
        i += 1

    # Save
    doc.save(output_path)
    print(f"✅ Resume generated: {output_path}")
    return True


def main():
    parser = argparse.ArgumentParser(
        description='Generate pixel-perfect resume from markdown (V3 - Template Match)',
        epilog='Example: python3 generate_resume_v3.py -i Template.md -o Template_REGENERATED.docx'
    )
    parser.add_argument('--input', '-i', required=True, help='Input RESUME.md file')
    parser.add_argument('--output', '-o', required=True, help='Output Resume.docx file')

    args = parser.parse_args()

    try:
        create_resume_from_markdown(args.input, args.output)
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == '__main__':
    exit(main())
